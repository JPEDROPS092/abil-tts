"""Extract supported documents into clean text suitable for TTS preparation."""
from __future__ import annotations

import os
import re
from html.parser import HTMLParser


class _HTMLTextExtractor(HTMLParser):
    _block_tags = {"article", "br", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li", "p", "section", "table", "tr"}
    _skip_tags = {"script", "style", "svg"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._skip_tags:
            self._skip_depth += 1
        if not self._skip_depth and tag in self._block_tags:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._skip_tags and self._skip_depth:
            self._skip_depth -= 1
        if not self._skip_depth and tag in self._block_tags:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return "".join(self.parts)


class _HTMLStructuredExtractor(HTMLParser):
    """Flatten HTML into markdown-like text, preserving document structure.

    Headings become ``#`` markers, ``<pre>`` becomes fenced code, tables become
    pipe rows and list items become ``-`` bullets so downstream classification
    can emit typed blocks.
    """

    _skip_tags = {"script", "style", "svg", "head", "noscript"}
    _block_tags = {
        "address", "article", "aside", "blockquote", "dd", "div", "dl", "dt",
        "fieldset", "figcaption", "figure", "footer", "form", "header", "hr",
        "main", "nav", "ol", "p", "section", "table", "ul",
    }
    _heading_tags = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._skip_depth = 0
        self._in_pre = False
        self._in_cell = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._skip_tags:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if self._in_pre:
            return
        if tag == "pre":
            self._in_pre = True
            self.parts.append("\n\n```\n")
        elif tag in self._heading_tags:
            self.parts.append("\n\n" + "#" * self._heading_tags[tag] + " ")
        elif tag == "li":
            self.parts.append("\n- ")
        elif tag == "tr":
            self.parts.append("\n|")
        elif tag in {"td", "th"}:
            self._in_cell = True
        elif tag == "br":
            self.parts.append("\n")
        elif tag in self._block_tags:
            self.parts.append("\n\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._skip_tags and self._skip_depth:
            self._skip_depth -= 1
            return
        if self._skip_depth or self._in_pre and tag != "pre":
            return
        if tag == "pre":
            self._in_pre = False
            self.parts.append("\n```\n\n")
        elif tag in {"td", "th"}:
            self._in_cell = False
            self.parts.append("|")
        elif tag in self._heading_tags or tag in self._block_tags:
            self.parts.append("\n\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._in_cell:
            self.parts.append(" " + " ".join(data.split()) + " ")
        else:
            self.parts.append(data)

    def text(self) -> str:
        return "".join(self.parts)


def _tidy_structure(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _clean_document_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?m)^\s*\d+\s*$", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_markdown(text: str) -> str:
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`[^`\n]+`", "", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*{1,3}([^*\n]+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_\n]+?)_{1,3}", r"\1", text)
    text = re.sub(r"!\[[^\]]*\]\([^\)]*\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"(?m)^\s*\|?\s*:?-{3,}:?(?:\s*\|\s*:?-{3,}:?)+\s*\|?\s*$", "", text)
    text = re.sub(r"(?m)^\s*\|\s*(.+?)\s*\|\s*$", lambda match: "; ".join(part.strip() for part in match.group(1).split("|")), text)
    return _clean_document_text(text)


def read_txt(path: str) -> str:
    with open(path, encoding="utf-8-sig", errors="replace") as handle:
        return _clean_document_text(handle.read())


def read_markdown(path: str) -> str:
    with open(path, encoding="utf-8-sig", errors="replace") as handle:
        return _strip_markdown(handle.read())


def read_html(path: str) -> str:
    with open(path, encoding="utf-8-sig", errors="replace") as handle:
        parser = _HTMLTextExtractor()
        parser.feed(handle.read())
    return _clean_document_text(parser.text())


def read_rtf(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as handle:
        text = handle.read()
    text = re.sub(r"\\par[d]?", "\n", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
    text = text.replace("{", "").replace("}", "").replace("\\", "")
    return _clean_document_text(text)


def read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx not installed. Run: pip install python-docx") from exc

    document = Document(path)
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("Tabela: " + "; ".join(cells) + ".")
    return _clean_document_text("\n\n".join(parts))


def read_pdf(path: str) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ImportError("PyMuPDF not installed. Run: pip install pymupdf") from exc

    document = fitz.open(path)
    try:
        pages = [page.get_text("text", sort=True) for page in document]
    finally:
        document.close()
    return _clean_document_text("\n\n".join(pages))


def read_epub(path: str) -> str:
    try:
        from ebooklib import ITEM_DOCUMENT, epub
    except ImportError as exc:
        raise ImportError("EbookLib not installed. Run: pip install EbookLib") from exc

    book = epub.read_epub(path)
    parts: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        parser = _HTMLTextExtractor()
        parser.feed(item.get_content().decode("utf-8", errors="replace"))
        content = _clean_document_text(parser.text())
        if content:
            parts.append(content)
    return _clean_document_text("\n\n".join(parts))


def _normalize_markdown(text: str) -> str:
    """Keep document structure (headings, fences, tables, lists); strip inline noise."""
    text = re.sub(r"!\[[^\]]*\]\([^\)]*\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    text = re.sub(r"\*{1,3}([^*\n]+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_\n]+?)_{1,3}", r"\1", text)
    text = re.sub(r"^>[ \t]?", "", text, flags=re.MULTILINE)
    text = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", text)
    return _tidy_structure(text)


def read_markdown_structured(path: str) -> str:
    with open(path, encoding="utf-8-sig", errors="replace") as handle:
        return _normalize_markdown(handle.read())


def read_html_structured(path: str) -> str:
    with open(path, encoding="utf-8-sig", errors="replace") as handle:
        parser = _HTMLStructuredExtractor()
        parser.feed(handle.read())
    return _tidy_structure(parser.text())


def read_epub_structured(path: str) -> str:
    try:
        from ebooklib import ITEM_DOCUMENT, epub
    except ImportError as exc:
        raise ImportError("EbookLib not installed. Run: pip install EbookLib") from exc

    book = epub.read_epub(path)
    parts: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        parser = _HTMLStructuredExtractor()
        parser.feed(item.get_content().decode("utf-8", errors="replace"))
        content = _tidy_structure(parser.text())
        if content:
            parts.append(content)
    return _tidy_structure("\n\n".join(parts))


READERS = {
    ".txt": read_txt,
    ".md": read_markdown,
    ".markdown": read_markdown,
    ".html": read_html,
    ".htm": read_html,
    ".rtf": read_rtf,
    ".docx": read_docx,
    ".pdf": read_pdf,
    ".epub": read_epub,
}

STRUCTURED_READERS = {
    ".md": read_markdown_structured,
    ".markdown": read_markdown_structured,
    ".html": read_html_structured,
    ".htm": read_html_structured,
    ".epub": read_epub_structured,
}


def read_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    reader = READERS.get(ext)
    if reader is None:
        supported = ", ".join(READERS)
        raise ValueError(f"Unsupported format '{ext}'. Supported: {supported}")
    return reader(path)


def read_document_structured(path: str) -> str:
    """Like read_document, but preserves structure (headings/code/tables) as light markdown."""
    ext = os.path.splitext(path)[1].lower()
    reader = STRUCTURED_READERS.get(ext) or READERS.get(ext)
    if reader is None:
        supported = ", ".join(READERS)
        raise ValueError(f"Unsupported format '{ext}'. Supported: {supported}")
    return reader(path)
